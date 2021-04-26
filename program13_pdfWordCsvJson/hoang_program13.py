# _____________________________________________________________________________ #
# Written by: Vy Hoang                                                          #
# This program demonstrates using Python to work with PDF and                   #
# MS Word documents, and CSV and JSON data.                                     #
# Use classes and functions to organize the functionality of this program.      #
# Classes are PDFProcessing, WordProcessing, CSVProcessing, and JSONProcessing. #
# _____________________________________________________________________________ #



import PyPDF2
import csv
import docx
import json
import pyinputplus as pypip

# Define class PDFProcessing to process the pdf file data
class PDFProcessing:
    def __init__(self, pdfFileName):
        self.pdfFile = pdfFileName

    # Return number of pages in pdf file
    def displayNumPages(self):
        pdfFileObj = open(self.pdfFile, 'rb')
        pdfNumPages = PyPDF2.PdfFileReader(pdfFileObj).numPages
        pdfFileObj.close()
        return pdfNumPages

    # Return text of the page
    def displayTextPage(self, pageNum):
        pdfFileObj = open(self.pdfFile, 'rb')
        pdfPage = PyPDF2.PdfFileReader(pdfFileObj).getPage(pageNum)
        pdfTextPage = pdfPage.extractText()
        pdfFileObj.close()
        return pdfTextPage

# Define WordProcessing class to process the .doc file data
class WordProcessing:
    def __init__(self, docFileName):
        self.docFile = docx.Document(docFileName)

    # Return total paragraphs of the file
    def displayNumParagraphs(self):
        return len(self.docFile.paragraphs)

    # return text of the paragraph
    def displayTextParagraph(self, paraNum):
        return self.docFile.paragraphs[paraNum].text

# Define CSVProcessing class to process the .csv file data
class CSVProcessing:
    def __init__(self, csvFilename):
        self.csvFile = csvFilename

    # Display csv file content
    def displayCsvFile(self):
        csvInFile = open(self.csvFile)
        csvReader = csv.reader(csvInFile)
        for row in csvReader:
            print('Row #', str(csvReader.line_num) + ' ' + str(row))
        csvInFile.close()

    # Update csv file content
    def updateCsvFile(self, newRow):
        csvOutFile = open(self.csvFile, 'a', newline='')
        csvWriter = csv.writer(csvOutFile)
        csvWriter.writerow(newRow)
        csvOutFile.close()


class JSONProcessing:
    def __init__(self, pythonDict):
        self.pythonDict = pythonDict

    # Convert Python dict data to JSON-formatted string
    def displayJsonData(self):
        stringJsonData = json.dumps(self.pythonDict, indent=4)
        return stringJsonData


def main():
    # Display number of pages in the pdf file
    pdfProcessor = PDFProcessing(pdfFileName='meetingminutes.pdf')
    print("\nThe number of pages in meetingminutes.pdf: ", pdfProcessor.displayNumPages())

    # Get page number and display the text on that page
    pageNum = int(input("\nEnter a page number: "))
    print("\nThe text on page #{}: \n".format(pageNum), pdfProcessor.displayTextPage(pageNum))

    # Display the number of paragraphs in the demo.docx
    wordProcecssor = WordProcessing('demo.docx')
    print("\nThe number of paragraphs in demo.docx: ", wordProcecssor.displayNumParagraphs())

    # Get paragraph number and display the text of the paragraph
    paraNum = int(input("\nEnter a paragraph number: "))
    print("\nThe text of the paragraph #{}: ".format(paraNum), wordProcecssor.displayTextParagraph(paraNum))

    # Display the content of example.csv
    csvFileProcessor = CSVProcessing(csvFilename='example.csv')
    print("\nThe contents of example.csv:")
    csvFileProcessor.displayCsvFile()

    # Get data input and update the data to the example.csv
    newRow = [input("\nEnter datetime: "),
              input("\nEnter fruit name: ").capitalize(),
              pypip.inputInt("\nEnter fruit quantity: ", blank=True)]
    csvFileProcessor.updateCsvFile(newRow=newRow)

    # Get dat input for 7 cities and store the dat in the dict
    cityDict = {}
    for num in range(7):
        city = input("\nEnter city name #{}: ".format(num + 1)).capitalize()
        cityDict[city] = input("\nEnter an adjective for {}: ".format(city)).capitalize()

    # Convert the Python dict to a string of JSON-formatted data and display JSON data
    jsonProcessor = JSONProcessing(pythonDict=cityDict)
    print("\nThe JSON data converted from Python dictionary:\n", jsonProcessor.displayJsonData())


if __name__ == '__main__':
    main()
else:
    pass

